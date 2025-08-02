import streamlit as st
from resume_generator import generate_resume_docx
from docx import Document
from io import BytesIO

st.set_page_config(page_title="TailorCV – AI Resume & Cover Letter Coach", layout="centered")
st.title("🧵 TailorCV – Build Smarter Resumes & Cover Letters")

option = st.radio("Choose what you want to do:",
    ["📤 Upload & Improve Resume", 
     "🧱 Build Resume from Scratch",
     "✍️ Generate a Cover Letter"])

# ------------------ Upload Resume ------------------
if option == "📤 Upload & Improve Resume":
    uploaded_file = st.file_uploader("Upload your resume (.docx or .pdf)", type=["docx", "pdf"])
    jd_text = st.text_area("Paste the Job Description here", placeholder="E.g., We're looking for a Data Analyst with experience in SQL, Python, and Power BI...")

    if uploaded_file and jd_text:
        st.success("Resume and JD uploaded successfully!")
        if st.button("Analyze Resume"):
            st.write("🧠 AI Feedback coming soon...")
            st.write("✍️ Cover Letter generation coming soon...")

# ------------------ Build Resume from Scratch ------------------
elif option == "🧱 Build Resume from Scratch":
    st.subheader("🔹 Resume Builder")

    demo_mode = st.checkbox("🧪 Use demo/test data")

    if demo_mode:
        name = "Pratiksha Raturi"
        phone = "9876543210"
        email = "pratiksha.raturi@example.com"
        linkedin = "https://linkedin.com/in/pratiksha-raturi"
        github = "https://github.com/pratiksha-raturi"
        summary = "Final-year B.Tech student with projects in resume AI and credit modeling. Skilled in Python, Streamlit, and ML."
        education = "B.Tech in CSE – UPES – 2021–2025 – CGPA: 8.3/10\nXII – DPS – 92%\nX – DPS – 94%"
        experience = "AI Intern | RANGR Data | May–Aug 2024\n• Built resume parser using OpenAI\n• Integrated LLM into analytics pipeline"
        projects = "TailorCV – Resume Builder | GPT + Streamlit\n• Built resume + cover letter tool with form inputs\n• Generates Word output using python-docx\n\nLandslide Predictor | CNN + HDF5 Data\n• Used image segmentation for terrain classification\n• Achieved 92% validation accuracy"
        skills = "Python, Streamlit, Pandas, Git, Power BI"
        certifications = "IBM GenAI – June 2025\n• Built a mini LLM project using prompts and embeddings\n• Created chatbot-like QA interface\n\nGoogle DA – April 2025\n• Used real-world datasets to build dashboards\n• Final project: bakery sales report in Google Sheets"
        submitted = True
    else:
        with st.form("resume_form"):
            name = st.text_input("Full Name")
            phone = st.text_input("Phone Number")
            email = st.text_input("Email Address")
            linkedin = st.text_input("LinkedIn URL", placeholder="https://linkedin.com/in/yourname")
            github = st.text_input("GitHub URL", placeholder="https://github.com/yourusername")

            summary = st.text_area("Professional Summary", placeholder="E.g., Final-year B.Tech student with internship experience in AI/ML and strong analytical background.", help="Summarize your career goals, key skills, and notable achievements.")

            education = st.text_area("Education", placeholder="E.g., B.Tech in Computer Science – UPES – 2025 – CGPA: 8.2\nXII – CBSE – 91%\nX – CBSE – 93%", help="Mention degree, institution, year, and scores in consistent format.")

            experience = st.text_area("Work Experience", placeholder="E.g., AI Intern | ABC Corp | Jan–Apr 2024\n• Built a chatbot using GPT-4\n• Conducted market research", help="Use bullet points with achievements and metrics.")

            projects = st.text_area("Projects (Optional)", placeholder="E.g.,\nTailorCV – Resume Builder | GPT + Streamlit\n• Built resume + cover letter tool with form inputs\n• Generates Word output using python-docx\n\nLandslide Predictor | CNN + HDF5 Data\n• Used image segmentation for terrain classification\n• Achieved 92% validation accuracy", help="Add academic or personal projects with tools used and 1–2 bullet points.")

            skills = st.text_area("Skills", placeholder="E.g., Python, SQL, Power BI, Streamlit, Git", help="Comma-separated list of key skills relevant to your role.")

            certifications = st.text_area("Certifications (Optional)", placeholder="E.g., IBM GenAI – June 2025\n• Built a mini LLM project\n• Created chatbot using prompts")

            submitted = st.form_submit_button("✅ Generate Resume")

    if submitted:
        data = {
            'name': name,
            'phone': phone,
            'email': email,
            'linkedin': linkedin,
            'github': github,
            'summary': summary,
            'education': education,
            'experience': experience,
            'projects': projects,
            'skills': skills,
            'certifications': certifications
        }

        docx_file = generate_resume_docx(data)
        st.success("✅ Resume generated successfully!")
        st.download_button("📥 Download Resume (.docx)", data=docx_file, file_name=f"{name.replace(' ', '_')}_Resume.docx")

# ------------------ Generate Cover Letter ------------------
elif option == "✍️ Generate a Cover Letter":
    st.subheader("✍️ Cover Letter Generator")

    demo_mode = st.checkbox("🧪 Use demo/test data")

    if demo_mode:
        your_name = "Pratiksha Raturi"
        your_institution = "UPES Dehradun"
        your_position = "Data Analyst Intern"
        company = "Zomato"
        experience_summary = "Exposure to LLM integrations, content writing, and business insights."
        resume_summary = "Experience in resume AI, ML pipelines, and prompt engineering with GPT-4."
        best_internship = "AI Intern at RANGR Data, built OpenAI-powered dashboards."
        how_will_help = "This opportunity aligns with my long-term analytics goals and will help me deepen my impact in research-based product teams."
        soft_skills = "Empathetic, collaborative, highly communicative with a passion for learning."
        submitted = True
    else:
        with st.form("cover_form"):
            your_name = st.text_input("Your Name")
            your_institution = st.text_input("College/University")
            your_position = st.text_input("Job Role You're Applying For")
            company = st.text_input("Company Name")
            experience_summary = st.text_area("Brief Experience Summary", help="Summarize your exposure to work, community, or projects.")
            resume_summary = st.text_area("Resume Summary Highlights", help="What does your resume show: key skills, achievements, tools?")
            best_internship = st.text_area("Your Best Internship", help="Mention company, duration, and work you did.")
            how_will_help = st.text_area("How this opportunity helps you", help="Explain how it supports your learning/career.")
            soft_skills = st.text_area("Mention your soft skills and work traits", help="E.g., collaborative, proactive, team-focused")
            submitted = st.form_submit_button("Generate Cover Letter")

    if submitted:
        doc = Document()
        doc.add_paragraph(your_name)
        doc.add_paragraph(your_institution)
        doc.add_paragraph("Complete Address | Contact Info")
        doc.add_paragraph("")
        doc.add_paragraph("Dear Selection Committee,")

        content = (
            f"Purpose: I am writing today in application for the position of {your_position} at {company}, India. "
            f"{experience_summary} "
            f"As my attached resume outlines, {resume_summary} "
            f"I wish to intensify my knowledge and follow this path in my long-term professional journey. "
            f"{how_will_help} "
            f"Apart from academics, my best internship experience was at {best_internship}. "
            f"{soft_skills} "
            f"\n\nYour careful review of my application is deeply appreciated. Thank you for the opportunity."
            f"\n\nSincerely,\n{your_name}"
        )
        doc.add_paragraph(content)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Cover Letter generated successfully!")
        st.download_button("📄 Download Cover Letter (.docx)", data=buffer, file_name=f"{your_name.replace(' ', '_')}_CoverLetter.docx")
