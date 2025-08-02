# resume_generator.py
from docx import Document
from io import BytesIO

def generate_resume_docx(data):
    doc = Document()
    doc.add_heading(data['name'], level=0)
    doc.add_paragraph(f"{data['phone']} | {data['email']} | {data['linkedin']} | {data['github']}")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(data['summary'])

    doc.add_heading("Education", level=1)
    doc.add_paragraph(data['education'])

    doc.add_heading("Experience / Projects", level=1)
    doc.add_paragraph(data['experience'])

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(data['skills'])

    if data['certifications']:
        doc.add_heading("Certifications", level=1)
        doc.add_paragraph(data['certifications'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
